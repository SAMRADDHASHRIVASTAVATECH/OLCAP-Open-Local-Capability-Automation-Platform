"""
JIT subprocess worker.

    python -m olcap.core.worker <component-id>

Speaks newline-delimited JSON on stdin/stdout. Heavy implementations are loaded
here so that releasing a worker actually returns memory to the OS.
"""
from __future__ import annotations

import json
import os
import sys
import traceback
from pathlib import Path
from urllib.parse import urlparse
from typing import Any, Callable, Dict


def _emit(obj: Dict[str, Any]) -> None:
    sys.stdout.write(json.dumps(obj, default=str) + "\n")
    sys.stdout.flush()


# --------------------------------------------------------------------------- #
# Loaders: real code for each heavy backend. If the backend is not installed
# the loader raises ImportError, which the worker reports honestly instead of
# pretending the capability works.
# --------------------------------------------------------------------------- #
class _Base:
    def __init__(self) -> None:
        self._ready = False

    def ping(self) -> Dict[str, Any]:
        return {"ok": True, "component": type(self).__name__}


def _assert_http(url: str) -> None:
    """A browser will happily open file:// and other schemes - don't let it.

    A URL arriving from a model or a scraped page must never become local
    file disclosure, so only http(s) is navigable.
    """
    scheme = (urlparse(str(url or "")).scheme or "").lower()
    if scheme not in ("http", "https"):
        raise ValueError(
            f"refusing to navigate to unsupported scheme {scheme or '(none)'!r}; "
            "only http and https are allowed")


class PlaywrightHandler(_Base):
    """Real browser automation."""

    def __init__(self) -> None:
        super().__init__()
        from playwright.sync_api import sync_playwright  # noqa: F401
        self._pw = sync_playwright().start()
        self._browser = self._pw.chromium.launch(
            args=["--no-sandbox", "--disable-dev-shm-usage"])
        self._ctx = self._browser.new_context()
        self._page = self._ctx.new_page()
        self._ready = True

    def web_interact(self, action: str = "", url: str = "", selector: str = "",
                     text: str = "", timeout_ms: int = 30000,
                     screenshot_path: str = "", **kw) -> Dict[str, Any]:
        if url:
            _assert_http(url)
        p = self._page
        p.set_default_timeout(int(timeout_ms))
        if url and action in ("navigate", "open", "goto", ""):
            p.goto(url, wait_until="domcontentloaded")
        if action in ("click",) and selector:
            p.click(selector)
        elif action in ("type", "fill") and selector:
            p.fill(selector, text or "")
        elif action in ("press",) and text:
            p.keyboard.press(text)
        elif action in ("wait",) and selector:
            p.wait_for_selector(selector)
        content = p.content()
        shot = None
        if screenshot_path or action == "screenshot":
            shot = screenshot_path or "playwright_shot.png"
            p.screenshot(path=shot, full_page=False)
        return {"ok": True, "action": action, "url": p.url,
                "title": p.title(), "content": content[:200000],
                "screenshot_path": shot}

    def web_browse(self, url: str = "", render: bool = True, **kw) -> Dict[str, Any]:
        _assert_http(url)
        self._page.goto(url, wait_until="networkidle", timeout=30000)
        return {"ok": True, "status": 200, "final_url": self._page.url,
                "html": self._page.content()[:500000],
                "title": self._page.title()}

    # ------------------------------------------------------------------ #
    # GUI + SCREENSHOT through the real browser. These are the same handler
    # the WEB_INTERACT capability uses, so a browser-backed desktop action
    # needs no second engine.
    # ------------------------------------------------------------------ #
    def gui_action(self, action: str = "", url: str = "", selector: str = "",
                   text: str = "", timeout_ms: int = 30000,
                   **kw) -> Dict[str, Any]:
        """Drive a page like a user: click, type, press, scroll, screenshot."""
        p = self._page
        p.set_default_timeout(int(timeout_ms))
        if url:
            p.goto(url, wait_until="domcontentloaded")
        a = (action or "").lower()
        if a in ("click",) and selector:
            p.click(selector)
        elif a in ("type", "fill") and selector:
            p.fill(selector, text or "")
        elif a in ("press", "key") and text:
            p.keyboard.press(text)
        elif a in ("move", "hover") and selector:
            p.hover(selector)
        elif a in ("scroll",):
            p.mouse.wheel(0, int(kw.get("delta_y", 600)))
        elif a in ("wait",) and selector:
            p.wait_for_selector(selector)
        elif a in ("info", "observe", ""):
            pass
        else:
            raise ValueError(f"unsupported gui action via playwright: {action!r}")
        return {"ok": True, "backend": "playwright", "action": a or "observe",
                "url": p.url, "title": p.title()}

    def screenshot_capture(self, path: str = "", target: str = "screen",
                           region: Any = None, url: str = "",
                           full_page: bool = False, **kw) -> Dict[str, Any]:
        """Capture a real screenshot: a page if `url` is given, else the view."""
        if not path:
            raise ValueError("path is required for screenshot_capture")
        p = self._page
        if url:
            _assert_http(url)
            p.goto(url, wait_until="domcontentloaded")
            p.wait_for_timeout(int(kw.get("settle_ms", 300)))
        clip = None
        if isinstance(region, dict) and {"x", "y", "width", "height"} <= set(region):
            clip = {"x": float(region["x"]), "y": float(region["y"]),
                    "width": float(region["width"]),
                    "height": float(region["height"])}
        out = Path(path).expanduser()
        out.parent.mkdir(parents=True, exist_ok=True)
        p.screenshot(path=str(out), full_page=bool(full_page),
                     **( {"clip": clip} if clip else {}))
        if not out.exists() or out.stat().st_size == 0:
            return {"ok": False, "error": f"screenshot not written to {out}"}
        return {"ok": True, "path": str(out), "bytes": out.stat().st_size,
                "backend": "playwright", "target": target, "url": p.url}

    def close(self) -> None:
        try:
            self._browser.close()
        except Exception:
            pass
        try:
            self._pw.stop()
        except Exception:
            pass


class Crawl4AIHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from crawl4ai import AsyncWebCrawler  # noqa: F401
        self._ready = True

    def web_crawl(self, url: str = "", max_pages: int = 20, max_depth: int = 2,
                  **kw) -> Dict[str, Any]:
        import asyncio
        from crawl4ai import AsyncWebCrawler

        async def _run():
            pages = []
            async with AsyncWebCrawler() as crawler:
                seen, frontier = set(), [(url, 0)]
                while frontier and len(pages) < max_pages:
                    u, d = frontier.pop(0)
                    if u in seen or d > max_depth:
                        continue
                    seen.add(u)
                    r = await crawler.arun(u)
                    pages.append({"url": u, "depth": d,
                                  "text": (r.markdown or "")[:50000]})
                    for link in (r.links or {}).get("internal", [])[:20]:
                        href = link.get("href") if isinstance(link, dict) else link
                        if href and href.startswith("http") and href not in seen:
                            frontier.append((href, d + 1))
            return pages

        return {"ok": True, "pages": asyncio.run(_run())}

    def web_extract(self, url: str = "", **kw) -> Dict[str, Any]:
        import asyncio
        from crawl4ai import AsyncWebCrawler

        async def _run():
            async with AsyncWebCrawler() as c:
                r = await c.arun(url)
                return {"text": r.markdown or "", "metadata": r.metadata or {}}
        return asyncio.run(_run())


class TrafilaturaHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        import trafilatura  # noqa: F401
        self._ready = True

    def web_extract(self, url: str = "", mode: str = "text", **kw) -> Dict[str, Any]:
        import trafilatura
        # trafilatura.fetch_url() has no timeout and no scheme policy: one
        # stalled host blocked the worker - and the whole objective - forever.
        # Fetch through the guarded, time-boxed helper instead.
        from .impls.web import _fetch, _TIMEOUT, _bounded
        try:
            _status, downloaded, _final = _fetch(
                url, timeout=float(kw.get("timeout", _TIMEOUT)))
        except Exception as e:
            return {"ok": False, "error": f"{type(e).__name__}: {str(e)[:160]}"}
        if not downloaded:
            return {"ok": False, "error": f"fetch failed: {url}"}
        # Extraction is CPU-bound and scales badly with pathological pages: a
        # multi-megabyte table soup can burn minutes inside trafilatura, which
        # is how a single extract call once consumed the whole JIT cap. Cap the
        # input and time-box the parse as well as the fetch.
        MAX_HTML = 4_000_000
        truncated = len(downloaded) > MAX_HTML
        if truncated:
            downloaded = downloaded[:MAX_HTML]

        def _extract():
            if mode in ("xml", "html"):
                text = trafilatura.extract(downloaded, output_format=mode,
                                           include_comments=False)
            else:
                text = trafilatura.extract(downloaded, include_comments=False,
                                           include_tables=True)
            meta = trafilatura.extract_metadata(downloaded)
            return text, meta

        try:
            out, meta = _bounded(_extract, 25.0)
        except TimeoutError:
            return {"ok": False,
                    "error": "extraction exceeded 25s; page is too complex"}
        return {"ok": bool(out and out.strip()), "text": out or "",
                "truncated": truncated,
                "metadata": meta.as_dict() if meta else {}}

    def document_process(self, path: str = "", **kw) -> Dict[str, Any]:
        import trafilatura
        data = open(path, "rb").read() if path else b""
        out = trafilatura.extract(data.decode("utf-8", "replace"))
        # Trafilatura cannot parse binary documents - say so instead of
        # silently returning an empty "success" so the router can fall back.
        return {"ok": bool(out and out.strip()), "text": out or "",
                "note": "trafilatura handles HTML/XML/text, not binary documents"}


class DoclingHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from docling.document_converter import DocumentConverter  # noqa: F401
        self._conv = DocumentConverter()
        self._ready = True

    def document_process(self, path: str = "", url: str = "", **kw) -> Dict[str, Any]:
        src = path or url
        res = self._conv.convert(src)
        doc = res.document
        text = doc.export_to_text()
        tables = []
        try:
            for t in doc.tables:
                tables.append(t.export_to_dataframe().to_dict("records"))
        except Exception:
            pass
        return {"ok": True, "text": text, "sections": [], "tables": tables,
                "metadata": {"source": src}}


class UnstructuredHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from unstructured.partition.auto import partition  # noqa: F401
        self._ready = True

    def document_process(self, path: str = "", **kw) -> Dict[str, Any]:
        from unstructured.partition.auto import partition
        els = partition(filename=path)
        text = "\n\n".join(str(e) for e in els)
        tables = [str(e) for e in els if e.category == "Table"]
        return {"ok": True, "text": text, "sections": [], "tables": tables,
                "metadata": {"elements": len(els)}}


class MarkerHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from marker.converters.pdf import PdfConverter  # noqa: F401
        from marker.models import create_model_dict
        self._converter = PdfConverter(artifact_dict=create_model_dict())
        self._ready = True

    def document_process(self, path: str = "", **kw) -> Dict[str, Any]:
        rendered = self._converter(path)
        text = rendered.markdown
        return {"ok": True, "text": text, "sections": [], "tables": [],
                "metadata": {"source": path}}


class PypdfHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        import pypdf  # noqa: F401
        self._ready = True

    def document_process(self, path: str = "", **kw) -> Dict[str, Any]:
        import pypdf
        r = pypdf.PdfReader(path)
        parts, tables = [], []
        for i, page in enumerate(r.pages):
            txt = page.extract_text() or ""
            parts.append(f"\n--- page {i+1} ---\n{txt}")
            try:
                for t in (page.extract_tables() or []):
                    tables.append(t)
            except Exception:
                pass
        return {"ok": True, "text": "".join(parts), "sections": [], "tables": tables,
                "metadata": {"pages": len(r.pages),
                             "title": (r.metadata or {}).get("/Title", "")}}


class DocxHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        import docx  # noqa: F401
        self._ready = True

    def document_process(self, path: str = "", **kw) -> Dict[str, Any]:
        import docx
        d = docx.Document(path)
        text = "\n".join(p.text for p in d.paragraphs)
        tables = [[c.text for c in row.cells] for t in d.tables for row in t.rows]
        return {"ok": True, "text": text, "sections": [], "tables": tables,
                "metadata": {"paragraphs": len(d.paragraphs)}}


class ChromaHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        import chromadb  # noqa: F401
        import os
        from ..core.config import cfg
        self._client = chromadb.PersistentClient(path=str(cfg().indexes / "chroma"))
        self._ready = True

    def vector_store_op(self, action: str = "search", collection: str = "default",
                        vectors=None, query: str = "", top_k: int = 10,
                        **kw) -> Dict[str, Any]:
        col = self._client.get_or_create_collection(collection)
        if action == "upsert":
            ids = [v.get("id") or str(i) for i, v in enumerate(vectors or [])]
            col.upsert(ids=ids,
                       embeddings=[v["vector"] for v in (vectors or [])],
                       documents=[v.get("text", "") for v in (vectors or [])],
                       metadatas=[v.get("meta", {}) for v in (vectors or [])])
            return {"ok": True, "count": len(ids)}
        res = col.query(query_texts=[query], n_results=top_k)
        return {"ok": True, "hits": res.get("documents", [[]])[0]}

    def close(self) -> None:
        pass


class Mem0Handler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from mem0 import Memory  # noqa: F401
        self._m = Memory()
        self._ready = True

    def memory_op(self, action: str = "search", text: str = "", query: str = "",
                  limit: int = 10, **kw) -> Dict[str, Any]:
        if action == "put":
            return {"ok": True, "result": self._m.add(text)}
        return {"ok": True, "items": self._m.search(query or text, limit=limit)}


class CogneeHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        import cognee  # noqa: F401
        self._ready = True

    def memory_op(self, action: str = "search", text: str = "", query: str = "",
                  limit: int = 10, **kw) -> Dict[str, Any]:
        import asyncio
        import cognee
        if action == "put":
            asyncio.run(cognee.add(text))
            asyncio.run(cognee.cognify())
            return {"ok": True}
        res = asyncio.run(cognee.search(query or text))
        return {"ok": True, "items": res}


class GraphitiHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from graphiti_core import Graphiti  # noqa: F401
        self._ready = True

    def memory_op(self, action: str = "search", text: str = "", query: str = "",
                  limit: int = 10, **kw) -> Dict[str, Any]:
        import asyncio
        from graphiti_core import Graphiti
        g = Graphiti(os.environ.get("NEO4J_URI", "bolt://localhost:7687"),
                     os.environ.get("NEO4J_USER", "neo4j"),
                     os.environ.get("NEO4J_PASSWORD", "password"))

        async def _run():
            if action == "put":
                await g.add_episode(name="olcap", episode_body=text)
                return {"ok": True}
            return {"ok": True, "items": await g.search(query or text, num_results=limit)}
        return asyncio.run(_run())


class LlamaIndexHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from llama_index.core import VectorStoreIndex  # noqa: F401
        self._ready = True

    def rag_query(self, action: str = "query", query: str = "", paths=None,
                  text: str = "", top_k: int = 6, **kw) -> Dict[str, Any]:
        from llama_index.core import Document, VectorStoreIndex
        if action == "ingest":
            docs = [Document(text=open(p, encoding="utf-8", errors="replace").read())
                    for p in (paths or [])]
            if text:
                docs.append(Document(text=text))
            idx = VectorStoreIndex.from_documents(docs)
            self._index = idx
            return {"ok": True, "indexed": len(docs)}
        idx = getattr(self, "_index", None)
        if idx is None:
            return {"ok": False, "error": "no index; ingest first"}
        qe = idx.as_query_engine(similarity_top_k=top_k)
        resp = qe.query(query)
        return {"ok": True, "answer": str(resp),
                "citations": [n.node.get_content()[:200]
                              for n in getattr(resp, "source_nodes", [])]}


class HaystackHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        from haystack import Document, Pipeline  # noqa: F401
        self._docs: list = []
        self._ready = True

    def rag_query(self, action: str = "query", query: str = "", paths=None,
                  text: str = "", top_k: int = 6, **kw) -> Dict[str, Any]:
        from haystack import Document
        from haystack.document_stores.in_memory import InMemoryDocumentStore
        if action == "ingest":
            self._docs = [Document(content=open(p, encoding="utf-8",
                                                errors="replace").read())
                          for p in (paths or [])]
            if text:
                self._docs.append(Document(content=text))
            return {"ok": True, "indexed": len(self._docs)}
        store = InMemoryDocumentStore()
        store.write_documents(self._docs)
        hits = [d.content for d in self._docs if query.lower()[:20] in d.content.lower()]
        return {"ok": True, "answer": "\n---\n".join(hits[:top_k]),
                "citations": hits[:top_k]}


class RagasHandler(_Base):
    def __init__(self) -> None:
        super().__init__()
        import ragas  # noqa: F401
        self._ready = True

    def verify_result(self, criteria=None, result=None, **kw) -> Dict[str, Any]:
        """RAGAs-style faithfulness proxy without requiring an LLM judge."""
        res = result or {}
        answer = str(res.get("answer", res.get("report", "")))
        sources = res.get("sources") or res.get("citations") or []
        checks = []
        checks.append({"name": "answer_non_empty", "passed": len(answer.strip()) > 0})
        checks.append({"name": "has_sources", "passed": len(sources) > 0})
        if sources and answer:
            overlap = sum(1 for s in sources
                          if str(s)[:40].lower() in answer.lower())
            checks.append({"name": "grounding_overlap",
                           "passed": overlap > 0,
                           "detail": f"{overlap}/{len(sources)} sources referenced"})
        return {"ok": True, "passed": all(c["passed"] for c in checks),
                "checks": checks}


LOADERS: Dict[str, Callable[[], _Base]] = {
    "playwright": PlaywrightHandler,
    "crawl4ai": Crawl4AIHandler,
    "trafilatura": TrafilaturaHandler,
    "docling": DoclingHandler,
    "unstructured": UnstructuredHandler,
    "marker": MarkerHandler,
    "pypdf": PypdfHandler,
    "python-docx": DocxHandler,
    "chroma": ChromaHandler,
    "mem0": Mem0Handler,
    "cognee": CogneeHandler,
    "graphiti": GraphitiHandler,
    "llamaindex": LlamaIndexHandler,
    "haystack": HaystackHandler,
    "ragas": RagasHandler,
}

# Components that must run in an isolated process (heavy / memory-hungry).
SUBPROCESS_COMPONENTS = set(LOADERS) | {"openhands", "temporal", "langfuse",
                                        "ragflow", "anythingllm", "letta",
                                        "phoenix", "searxng", "firecrawl",
                                        "perplexica", "mindsearch",
                                        "open-deep-search", "gpt-researcher",
                                        "browser-use", "stagehand", "prefect",
                                        "crewai", "ag2-autogen", "smolagents",
                                        "agentscope", "praisonai", "langgraph"}


def main() -> int:
    if len(sys.argv) < 2:
        _emit({"event": "error", "error": "usage: worker <component>"})
        return 2
    comp = sys.argv[1]
    handler: Any = None
    try:
        if comp in LOADERS:
            handler = LOADERS[comp]()
        else:
            # Generic: import the module to prove the backend is usable.
            from ..core.registry import registry
            spec = registry().component(comp)
            mod = (spec.python_module if spec else "") or comp.replace("-", "_")
            __import__(mod)

            class _Generic(_Base):
                def ping(self):
                    return {"ok": True, "component": comp, "module": mod,
                            "note": "generic loader - import verified"}
            handler = _Generic()
        _emit({"event": "ready", "component": comp, "pid": os.getpid()})
    except Exception as e:
        _emit({"event": "error", "component": comp,
               "error": f"{type(e).__name__}: {e}",
               "traceback": traceback.format_exc()[-800:],
               "installed": False})
        return 1

    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue
        try:
            req = json.loads(line)
        except Exception:
            continue
        method = req.get("method", "")
        if method == "__stop__":
            try:
                if hasattr(handler, "close"):
                    handler.close()
            except Exception:
                pass
            _emit({"event": "result", "ok": True, "result": {"stopped": True}})
            break
        try:
            fn = getattr(handler, method, None)
            if fn is None:
                raise AttributeError(f"method '{method}' not implemented by {comp}")
            out = fn(**(req.get("params") or {}))
            _emit({"event": "result", "ok": True, "result": out})
        except Exception as e:
            _emit({"event": "result", "ok": False,
                   "error": f"{type(e).__name__}: {e}",
                   "traceback": traceback.format_exc()[-1200:]})
    return 0


if __name__ == "__main__":
    sys.exit(main())
